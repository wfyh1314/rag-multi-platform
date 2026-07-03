"""docx、md、txt 纯文本解析."""

from pathlib import Path

from core.exceptions import FileParseError
from document.loaders.base import BaseLoader, Document


class OfficeLoader(BaseLoader):
    """解析 docx、txt、md 文档."""

    extensions = frozenset({".docx", ".txt", ".md", ".doc"})

    def load(self, file_path: Path) -> list[Document]:
        if not file_path.exists():
            raise FileParseError("文件不存在", filename=file_path.name)

        ext = file_path.suffix.lower()
        if ext == ".doc":
            raise FileParseError("不支持 .doc 格式，请转换为 .docx", filename=file_path.name)

        if ext == ".docx":
            text = self._load_docx(file_path)
        else:
            text = self._load_text(file_path)

        text = text.strip()
        if not text:
            raise FileParseError("文档内容为空", filename=file_path.name)

        return [
            Document(
                content=text,
                metadata={
                    "source": file_path.name,
                    "file_type": ext.lstrip("."),
                },
            )
        ]

    def _load_docx(self, path: Path) -> str:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise FileParseError("缺少 python-docx 依赖", filename=path.name) from exc

        try:
            doc = DocxDocument(str(path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as exc:
            raise FileParseError(f"docx 解析失败: {exc}", filename=path.name) from exc

    def _load_text(self, path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            try:
                return path.read_text(encoding="gbk")
            except Exception as exc:
                raise FileParseError(f"文本编码解析失败: {exc}", filename=path.name) from exc
